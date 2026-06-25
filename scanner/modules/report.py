"""
report.py
---------
Generates the final vulnerability assessment report.

This module takes all the data gathered during the assessment (recon,
Nessus findings, validation results, logging review) and assembles it into
a professional report following the same 8-section structure used in the
manual assessment reports.

It produces two formats:
  - a plain-text (.txt) report, easy to read and store in the repo
  - a Word (.docx) report, suitable for formal submission

The .docx generation uses the python-docx library.
"""

import os
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ---------------------------------------------------------------------------
# TEXT REPORT
# ---------------------------------------------------------------------------

def build_report_sections(config, recon_data, nessus_data,
                          validation_data, logging_data):
    """
    Build the eight report sections as a list of (title, body) pairs.

    Both the text and Word report builders use this same structure, so the
    two formats always stay consistent. Each 'body' is a plain string.

    This is the single source of truth for the report's content.
    """
    target_ip = config["target_ip"]
    date_str = datetime.now().strftime("%B %d, %Y")

    sections = []

    # --- Section 1: Scope and Objectives ---
    sections.append((
        "1. Scope and Objectives",
        "The purpose of this vulnerability assessment is to evaluate the "
        f"security posture of the system at {target_ip}. The assessment "
        "identifies security weaknesses, assesses their risk levels, and "
        "provides remediation strategies aligned with the CIS Critical "
        "Security Controls (CIS Controls v8.1) framework.\n\n"
        "The assessment covers open ports, running services, authentication "
        "mechanisms, encryption standards, firewall configuration, and "
        "logging capabilities. Tools used include Nmap for reconnaissance, "
        "Nessus for automated vulnerability scanning, and manual validation "
        "methods to confirm exploitability. All testing was conducted within "
        "a controlled environment for ethical and legal compliance."
    ))

    # --- Section 2: Asset Identification ---
    if recon_data:
        body = (
            f"An Nmap scan of {target_ip} identified the operating system as "
            f"{recon_data['os_guess']} and discovered "
            f"{len(recon_data['services'])} open ports.\n\n"
        )
        # Summarize each risk category
        for level in ["Critical", "High", "Medium", "Low"]:
            items = recon_data["categorized"][level]
            if not items:
                continue
            body += f"\n{level} Risk Assets:\n"
            for item in items:
                body += (f"  - Port {item['port']} ({item['service']}, "
                        f"{item['version']}): {item['reason']}\n")
    else:
        body = "Reconnaissance data was not available for this assessment."

    sections.append(("2. Asset Identification and Categorization", body))

    # --- Section 3: Security Controls Evaluation ---
    sections.append((
        "3. Security Controls and Policies Evaluation",
        "The evaluation examined access controls, encryption standards, "
        "network security, and patch management on the target.\n\n"
        "Findings from the automated scan and manual validation indicate the "
        "presence of weak or default credentials, outdated services, and "
        "missing network protections. Detailed findings appear in the "
        "following sections. Where default credentials or anonymous access "
        "were confirmed, these represent failures of access control. Where "
        "outdated services were found, these represent gaps in patch "
        "management that leave known vulnerabilities exposed."
    ))

    # --- Section 4: Threat Modeling and Risk Assessment ---
    if nessus_data:
        findings = nessus_data["findings"]
        total = sum(len(findings[lvl]) for lvl in findings)
        body = (
            f"A Nessus vulnerability scan returned {total} findings: "
            f"{len(findings['Critical'])} critical, "
            f"{len(findings['High'])} high, "
            f"{len(findings['Medium'])} medium, "
            f"{len(findings['Low'])} low, and "
            f"{len(findings['Info'])} informational.\n\n"
        )
        # List the critical and high findings by name
        for level in ["Critical", "High"]:
            items = findings[level]
            if not items:
                continue
            body += f"\n{level} Vulnerabilities:\n"
            for item in items:
                body += f"  - {item['name']} ({item['family']})\n"
    else:
        body = "Nessus scan data was not available for this assessment."

    sections.append(("4. Threat Modeling and Risk Assessment", body))

    # --- Section 5: Security Posture Testing and Validation ---
    if validation_data:
        body = ("Manual validation was performed to confirm the "
                "exploitability of key findings.\n\n")
        for r in validation_data["results"]:
            status = "CONFIRMED" if r["success"] else "Not confirmed"
            body += f"[{status}] {r['check']}\n    {r['detail']}\n\n"
    else:
        body = "Validation data was not available for this assessment."

    sections.append(("5. Security Posture Testing and Validation", body))

    # --- Section 6: Monitoring, Logging, and Incident Response ---
    if logging_data and logging_data.get("summary_text"):
        body = logging_data["summary_text"]
    else:
        body = ("Logging and firewall review could not be completed, often "
                "because SSH access was unavailable. Where review is not "
                "possible, the organization should manually verify that "
                "logging, alerting, and firewall controls are in place.")

    sections.append(("6. Monitoring, Logging, and Incident Response Readiness", body))

    # --- Section 7: Recommendations ---
    sections.append((
        "7. Recommendations and Security Hardening Measures",
        "The following measures are recommended, aligned with the CIS "
        "Critical Security Controls:\n\n"
        "  - Remove or replace outdated and backdoored services such as "
        "vsftpd (CIS Control 2: Inventory and Control of Software Assets).\n"
        "  - Change all default credentials and enforce a strong password "
        "policy (CIS Control 5: Account Management).\n"
        "  - Restrict NFS exports and other file-sharing services to trusted "
        "hosts only (CIS Control 4: Secure Configuration of Assets).\n"
        "  - Disable legacy plaintext protocols (Telnet, rlogin) and enforce "
        "SSH key-based authentication (CIS Control 6: Access Control "
        "Management).\n"
        "  - Update all outdated software and the operating system itself "
        "(CIS Control 7: Continuous Vulnerability Management).\n"
        "  - Implement firewall rules to restrict access to required ports "
        "only (CIS Control 12: Network Infrastructure Management).\n"
        "  - Deploy centralized logging and a SIEM for threat detection "
        "(CIS Control 8: Audit Log Management).\n\n"
        "Applying these controls would significantly reduce the attack "
        "surface and improve the system's resilience."
    ))

    # --- Section 8: Summary and Conclusion ---
    # Build a short data-driven summary line
    if nessus_data:
        findings = nessus_data["findings"]
        crit = len(findings["Critical"])
        high = len(findings["High"])
        summary_line = (f"The assessment identified {crit} critical and "
                       f"{high} high severity vulnerabilities ")
    else:
        summary_line = "The assessment identified multiple vulnerabilities "

    sections.append((
        "8. Summary and Conclusion",
        f"{summary_line}on the target system at {target_ip}. "
        "The most significant risks stem from weak or default authentication, "
        "outdated software, and missing network protections, all of which "
        "could allow an attacker to gain unauthorized access.\n\n"
        "Immediate remediation should focus on removing backdoored services, "
        "changing default credentials, restricting exposed file shares, and "
        "implementing firewall rules. Without these measures, the system "
        "remains at high risk of compromise.\n\n"
        "Regular vulnerability assessments aligned to the CIS Critical "
        "Security Controls should be conducted to ensure continuous "
        f"improvement of the system's security posture.\n\n"
        f"Report generated: {date_str}"
    ))

    return sections


def generate_txt_report(sections, output_path, target_ip):
    """
    Write the report sections to a plain-text file.
    """
    print("[*] Generating text report ...")

    lines = []
    lines.append("VULNERABILITY ASSESSMENT REPORT")
    lines.append("=" * 55)
    lines.append(f"Target System: {target_ip}")
    lines.append(f"Date: {datetime.now().strftime('%B %d, %Y')}")
    lines.append(f"Framework: CIS Critical Security Controls v8.1")
    lines.append("=" * 55)
    lines.append("")

    # Add each section in order
    for title, body in sections:
        lines.append("")
        lines.append(title)
        lines.append("-" * 55)
        lines.append(body)
        lines.append("")

    with open(output_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))

    print(f"[+] Text report saved: {output_path}")
    return output_path


# ---------------------------------------------------------------------------
# WORD (DOCX) REPORT
# ---------------------------------------------------------------------------

def generate_docx_report(sections, output_path, target_ip):
    """
    Write the report sections to a Word (.docx) document.

    Uses python-docx to build a clean, professional document with a title
    page, headings, and body text — matching the simple prose style of the
    manual reports.
    """
    print("[*] Generating Word report ...")

    doc = Document()

    # --- Set a default font for the whole document ---
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    # --- Title ---
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Vulnerability Assessment Report")
    run.bold = True
    run.font.size = Pt(20)

    # --- Subtitle / metadata ---
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Target System: {target_ip}\n").font.size = Pt(11)
    meta.add_run(
        f"Date: {datetime.now().strftime('%B %d, %Y')}\n"
    ).font.size = Pt(11)
    meta.add_run(
        "Framework: CIS Critical Security Controls v8.1"
    ).font.size = Pt(11)

    doc.add_paragraph()  # spacer

    # --- Each section ---
    for title_text, body_text in sections:
        # Section heading
        heading = doc.add_paragraph()
        run = heading.add_run(title_text)
        run.bold = True
        run.font.size = Pt(14)

        # Section body — split on newlines so it lays out nicely
        for paragraph_text in body_text.split("\n"):
            if paragraph_text.strip() == "":
                continue
            p = doc.add_paragraph(paragraph_text)
            p.paragraph_format.space_after = Pt(6)

        doc.add_paragraph()  # spacer between sections

    # --- Disclaimer footer ---
    disclaimer = doc.add_paragraph()
    disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = disclaimer.add_run(
        "This report was produced for educational purposes in a controlled "
        "lab environment."
    )
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.save(output_path)
    print(f"[+] Word report saved: {output_path}")
    return output_path


def generate_reports(config, recon_data, nessus_data,
                    validation_data, logging_data, output_dir):
    """
    The main entry point for this module. Builds the report sections once,
    then writes them out in both text and Word formats.

    Returns a tuple of the two output file paths.
    """
    print("\n[*] Generating final reports ...")

    target_ip = config["target_ip"]
    assessment_id = config["assessment_id"]

    # Build the section content once, shared by both formats
    sections = build_report_sections(
        config, recon_data, nessus_data, validation_data, logging_data
    )

    # Make sure the output directory exists
    os.makedirs(output_dir, exist_ok=True)

    # Build file paths
    txt_path = os.path.join(output_dir, f"report_{assessment_id}.txt")
    docx_path = os.path.join(output_dir, f"report_{assessment_id}.docx")

    # Generate both formats
    generate_txt_report(sections, txt_path, target_ip)
    generate_docx_report(sections, docx_path, target_ip)

    return txt_path, docx_path


# Allows testing this module on its own with placeholder data:
#   python modules/report.py
if __name__ == "__main__":
    print("Testing report module with sample data ...")

    sample_config = {
        "target_ip": "192.168.56.103",
        "assessment_id": "test_demo",
    }

    # Minimal fake data just to prove the report builds
    sample_recon = {
        "os_guess": "Linux 2.6.x (Ubuntu 8.04)",
        "services": [{"port": 21, "service": "ftp", "version": "vsftpd 2.3.4"}],
        "categorized": {
            "Critical": [{"port": 21, "service": "ftp",
                         "version": "vsftpd 2.3.4",
                         "reason": "Known backdoor"}],
            "High": [], "Medium": [], "Low": []
        }
    }

    txt, docx = generate_reports(
        sample_config, sample_recon, None, None, None, "reports"
    )
    print(f"\nDone. Created:\n  {txt}\n  {docx}")